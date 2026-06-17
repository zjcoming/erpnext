from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"E:\projects\vscode\erpnext\outputs\ERPNext制造与计件工资一期合作范围确认建议稿.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
GREEN = "E7F4EA"
AMBER = "FFF4D6"
RED = "FCE8E6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, name="Calibri", east_asia="Microsoft YaHei", size=11, color="000000"):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)


def add_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_nums = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract or [0]) + 1
    num_id = max(existing_nums or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, num_id, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        set_run_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic)
    else:
        set_run_font(p.add_run(text), italic=italic)
    return p


def add_callout(doc, title, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(title), size=11, bold=True, color=INK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    set_run_font(p2.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, status_colors=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(header), size=9.5, bold=True, color=INK)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if idx == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(str(value)), size=font_size)
            if status_colors and idx == status_colors[0]:
                fill = status_colors[1].get(str(value))
                if fill:
                    set_cell_shading(cells[idx], fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
set_style_font(normal, size=11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 14, 8),
    ("Heading 2", 13, BLUE, 11, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    set_style_font(style, size=size, color=color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run_font(hp.add_run("ERPNext 制造执行与计件工资项目"), size=9, color=MID_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(fp.add_run("一期合作范围确认建议稿 | 讨论版"), size=8.5, color=MID_GRAY)

num_id = add_numbering_definition(doc)

# Cover / customer pack
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after = Pt(3)
set_run_font(p.add_run("合作范围确认建议稿"), size=11, bold=True, color=BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
set_run_font(p.add_run("ERPNext 制造执行与计件工资系统"), size=28, bold=True, color=INK)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(24)
set_run_font(p.add_run("以“缩短订单后操作链路、让工人容易报工、自动形成计件工资”为一期核心"), size=13, color=MID_GRAY)

meta = doc.add_table(rows=4, cols=2)
set_table_geometry(meta, [2100, 7260])
meta.style = "Table Grid"
metadata = [
    ("文档状态", "讨论版，用于确认合作内容，不等同于最终合同"),
    ("建议周期", "8周（以范围冻结、基础资料按时提供为前提）"),
    ("核心范围", "生产制造流程简化、移动报工、质量结果、个人计件工资、基础报表"),
    ("编制日期", "2026年6月14日"),
]
for row, (label, value) in zip(meta.rows, metadata):
    set_cell_shading(row.cells[0], LIGHT_GRAY)
    set_run_font(row.cells[0].paragraphs[0].add_run(label), size=10, bold=True, color=INK)
    set_run_font(row.cells[1].paragraphs[0].add_run(value), size=10)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_callout(
    doc,
    "本稿的核心判断",
    "客户原始需求是一份完整ERP需求池，覆盖进销存、制造、质量、计件、考勤、工资、财务、APP/小程序和离线同步。"
    "若全部纳入一期，两个月无法可靠交付。建议一期只承诺制造主链路和个人计件工资闭环，其余作为二期需求池单独评估。",
    AMBER,
)

add_heading(doc, "1. 项目目标与成功标准", 1)
add_body(doc, "一期不是重做一套ERP，而是在ERPNext现有能力上，把客户最常用的制造流程压缩成更容易操作的业务路径。")
for text in [
    "销售人员提交订单后，系统自动计算库存、生产需求和物料缺口，并生成待确认的生产/采购建议。",
    "计划员在一个“生产指挥台”内完成确认，不必反复进入多个标准单据页面。",
    "工人手机端只看到本人可开工的任务，执行开始、暂停、合格、不良、报废、返工等最少操作。",
    "质检确认后的有效产量自动进入个人计件明细，月底形成可审核的工资计算结果。",
    "保留ERP必须的审核、库存、成本和操作日志，不通过简单隐藏流程破坏数据准确性。",
]:
    add_bullet(doc, num_id, text)

add_heading(doc, "2. 建议的一期业务流程", 1)
flow_rows = [
    ("1", "销售订单", "录入客户、产品、数量、交期并提交", "销售"),
    ("2", "需求与齐套核验", "自动展开BOM，核验成品库存、在途、原料缺口和预计交期", "系统"),
    ("3", "生产确认", "一键生成生产工单、工序任务和缺料建议；关键单据仍由负责人确认提交", "计划员"),
    ("4", "采购与到货", "缺料形成采购物料请求，采购审核后生成采购订单；到货、质检、入库", "采购/仓库"),
    ("5", "任务释放", "材料齐套且前置工序满足条件后，任务显示为可开工", "系统/班组长"),
    ("6", "移动报工", "工人开始、暂停、分批报工，录入合格、不良、报废和返工数量", "工人"),
    ("7", "质量确认", "按规则抽检或全检，确认有效产量、返工和报废结果", "质检"),
    ("8", "完工与计件", "完工入库；有效产量按产品+工序单价形成计件金额", "系统"),
    ("9", "月度结算", "生成个人计件月报和工资计算草稿，经人事、财务审批", "人事/财务"),
]
add_table(doc, ["步骤", "环节", "一期目标操作", "责任角色"], flow_rows, [620, 1500, 5580, 1660], font_size=9)
add_callout(
    doc,
    "自动化边界",
    "系统可以自动计算并生成建议或草稿，但采购订单、生产工单、工资结果等关键业务单据建议保留人工确认。"
    "这既能减少操作，也能避免错误BOM、错误数量或异常交期被系统直接执行。",
)

add_page_break(doc)
add_heading(doc, "3. 一期范围（建议作为合同验收范围）", 1)
scope_rows = [
    ("P0", "基础档案", "物料、BOM、工艺路线、工序、工作站、仓库、员工、班组、客户、供应商基础配置；提供导入模板。", "纳入"),
    ("P0", "角色化简化界面", "销售、计划、仓库、班组长、工人、质检、人事/财务分别显示必要菜单和快捷入口。", "纳入"),
    ("P0", "订单到生产快捷入口", "销售订单提交后展示库存、在途、已占用、生产缺口、原料缺口；确认后生成对应生产单据。", "纳入"),
    ("P0", "物料齐套与开工提示", "显示缺料、部分到货、已齐套、允许开工；工人默认只看到已释放任务。", "纳入"),
    ("P0", "工单与工序派工", "按工单自动生成工序任务；支持分配员工/班组、计划数量、计划时间、暂停和终止。", "纳入"),
    ("P0", "手机端工序报工", "响应式网页/PWA入口；支持开始、暂停、分批完成、合格、不良、报废、返工、原因和备注。", "纳入"),
    ("P0", "基础质量联动", "工序/完工质检；质检结果影响有效计件数量；保留修改与审核日志。", "纳入"),
    ("P0", "个人计件规则", "按产品+工序设置计件单价和生效日期；按审核后的有效产量计算个人计件金额。", "纳入"),
    ("P0", "月度计件结算", "计件日报、月报、个人明细、班组汇总；形成工资计算草稿并走审核。", "纳入"),
    ("P1", "管理看板", "订单、工单、工序、缺料、齐套、在制、完工、不良率和个人产量基础看板。", "纳入，控制复杂度"),
    ("P1", "打印与导出", "一期约定报表支持Excel导出；约定的工单、任务单、计件明细支持打印。", "纳入"),
    ("P1", "部署与培训", "部署、管理员培训、关键岗位培训、基础操作手册、备份策略。", "纳入"),
]
add_table(
    doc,
    ["优先级", "功能域", "一期交付内容", "结论"],
    scope_rows,
    [820, 1600, 5700, 1240],
    status_colors=(3, {"纳入": GREEN, "纳入，控制复杂度": AMBER}),
    font_size=8.7,
)

add_heading(doc, "4. 订单后操作流程的具体简化方案", 1)
for text in [
    "新增“订单生产确认”页面：集中显示销售订单、成品现货、已占用库存、待生产数量、BOM、物料缺口、在途采购和建议开工日期。",
    "提供“一键生成生产任务包”：生成生产工单和工序任务；缺购物料生成采购物料请求草稿；缺自制半成品生成下级生产建议。",
    "默认隐藏不常用字段和高级功能，标准单据仍保留给管理员处理异常情况。",
    "通过颜色和状态表达业务：红色缺料、黄色部分齐套、绿色可开工；工人无需理解预计库存、Bin、Stock Entry等系统术语。",
    "允许计划员选择自动化策略：仅生成草稿、生成并提交生产工单、采购保持审批；策略由客户在上线前确认。",
    "对重复性高的动作提供批量按钮，例如批量派工、批量释放、批量打印任务卡和批量生成物料请求。",
]:
    add_bullet(doc, num_id, text)

add_heading(doc, "5. 车间工人手机端", 1)
worker_rows = [
    ("我的任务", "只显示本人/本班组任务；按待开工、进行中、待质检、已完成分类。"),
    ("任务详情", "产品、工单、工序、计划数量、已完成、待完成、作业要求、物料状态。"),
    ("报工操作", "开始、暂停、继续、提交本次数量；支持多次阶段性报工。"),
    ("质量数量", "合格、不良、报废、返工；异常必须选择原因，可填写备注。"),
    ("管控", "不可超过可报数量；上道工序/齐套/质检限制按配置启用。"),
    ("个人收益", "查看本人当日、当月已审核计件数量和计件金额，不显示他人工资。"),
]
add_table(doc, ["页面/能力", "一期表现"], worker_rows, [2000, 7360], font_size=9.3)
add_callout(
    doc,
    "移动端交付形态",
    "一期建议交付手机浏览器/PWA快捷方式，而不是原生APP或微信小程序。当前ERPNext响应式页面已可在iPhone和Android使用，"
    "通过定制首页和报工页面即可满足试运行。原生APP、小程序和离线同步属于独立工程，不应默认包含在8周范围内。",
    AMBER,
)

add_page_break(doc)
add_heading(doc, "6. 个人计件工资一期规则", 1)
add_body(doc, "一期先完成“可解释、可追溯、可审核”的个人计件闭环，不一次性承诺所有复杂薪资模式。")

rate_rows = [
    ("计件单价", "产品 + 工序 + 单价 + 生效日期；历史单价保留，不覆盖已结算期间。"),
    ("有效数量", "默认以质检确认的合格数量计算；未质检、待复检数量暂不计薪。"),
    ("返工", "可配置为不计薪或按指定比例计薪；返工责任归属和重复计件规则必须上线前确认。"),
    ("报废/不良", "默认不计入有效数量；如需扣款，必须形成可审核的调整记录并保留原因。"),
    ("系数", "一期可保留单一工序/员工系数；复杂阶梯、加班、难度组合规则列入二期。"),
    ("人工调整", "补计、扣减必须填写原因和附件，经过权限审批，不允许直接改最终金额。"),
    ("月度结算", "按结算期间汇总，锁定已审核数据，生成个人计件明细和工资草稿。"),
]
add_table(doc, ["规则项", "一期建议"], rate_rows, [1900, 7460], font_size=9.2)

add_callout(
    doc,
    "一期计件公式",
    "个人计件金额 = Σ（已审核有效数量 × 产品工序单价 × 适用系数）+ 已审批补计 - 已审批扣减。"
    "月度实发工资可进一步由“基本工资 + 计件工资 + 补贴 - 考勤扣款 - 其他扣款”构成，"
    "但完整考勤设备对接和复杂薪资制度需另行确认。",
    GREEN,
)

add_heading(doc, "7. 一期暂不承诺或需单独报价的内容", 1)
excluded_rows = [
    ("原生APP、微信小程序", "一期用响应式网页/PWA；如必须上架或接入微信生态，单独立项。"),
    ("离线报工与断网同步", "涉及本地数据库、冲突合并和安全机制，技术风险高，二期专项。"),
    ("班组集体计件、阶梯计件、计时计件混合", "一期仅个人计件基础模型；复杂规则待真实样例确认后扩展。"),
    ("考勤机/第三方考勤自动对接", "一期可导入月度考勤结果；硬件/API对接单独评估。"),
    ("完整财务重构与税务接口", "ERPNext标准财务可使用，但本期不承诺中国本地财税深度适配。"),
    ("全量采购、销售、委外深度定制", "保留ERPNext标准能力；仅改造与制造主流程直接相关的入口和联动。"),
    ("所有报表任意拖拽设计", "一期交付约定报表；新增复杂报表按变更处理。"),
    ("全量白标和去除所有ERPNext痕迹", "可做基础Logo、色彩和名称；彻底白标、升级兼容和许可证义务另行确认。"),
    ("硬件扫码枪、电子秤、打印机适配", "标准键盘式扫码可测试；特定硬件驱动和模板单独验收。"),
]
add_table(doc, ["事项", "一期边界"], excluded_rows, [3000, 6360], font_size=9)

add_heading(doc, "8. 八周实施计划", 1)
milestone_rows = [
    ("第1周", "调研与范围冻结", "现场流程、角色、样例单据；确认一期清单、计件规则、验收案例。"),
    ("第2周", "原型与基础配置", "简化工作区、角色菜单、订单生产确认原型；导入基础资料样例。"),
    ("第3-4周", "制造主流程", "需求核验、生产任务包、缺料/齐套、派工、手机报工、工序流转。"),
    ("第5周", "质量与计件", "质量结果、计件单价、有效数量、个人计件明细和月度汇总。"),
    ("第6周", "看板与报表", "生产看板、缺料/进度/不良/计件报表、打印和导出。"),
    ("第7周", "客户试运行/UAT", "选1-2个真实产品走完整流程；培训关键用户；记录问题。"),
    ("第8周", "修复与上线", "修复阻断问题、数据确认、操作手册、备份、正式环境上线。"),
]
add_table(doc, ["周期", "阶段", "主要产出"], milestone_rows, [1200, 1900, 6260], font_size=9.2)
add_callout(
    doc,
    "进度前提",
    "客户需在第1周提供基础数据和规则负责人，并在每个里程碑后1-2个工作日内反馈。"
    "若基础资料、计件规则或验收人员长期未确认，交付日期应相应顺延。",
    AMBER,
)

add_page_break(doc)
add_heading(doc, "9. 双方责任与交付物", 1)
responsibility_rows = [
    ("实施方", "完成约定配置和定制开发；维护测试环境；提供导入模板；组织演示、培训和UAT支持；交付约定文档与代码。"),
    ("客户方", "指定业务负责人；提供真实BOM、工艺、单价、员工、仓库、异常原因和样例订单；及时确认流程和测试结果。"),
    ("共同", "第1周冻结一期范围；共同确定验收数据；对新增需求进行影响评估，不在口头沟通中自动扩大范围。"),
]
add_table(doc, ["责任方", "主要责任"], responsibility_rows, [1600, 7760], font_size=9.3)

deliverables = [
    "可运行的ERPNext定制系统和一期自定义App源代码。",
    "生产制造、移动报工、质量和个人计件工资的一期功能。",
    "基础数据导入模板及一轮样例数据导入支持。",
    "约定的生产与计件报表、打印模板。",
    "部署说明、备份说明、管理员手册和关键岗位操作手册。",
    "关键用户培训及一次完整UAT问题修复周期。",
]
add_heading(doc, "9.1 建议交付物清单", 2)
for text in deliverables:
    add_bullet(doc, num_id, text)

add_heading(doc, "10. 验收建议", 1)
acceptance_rows = [
    ("订单转生产", "新建销售订单后能显示库存和物料缺口，并按确认结果生成正确数量的生产单据。"),
    ("缺料与齐套", "缺料时任务不可误判为可开工；采购入库后齐套状态可更新。"),
    ("工序执行", "工人手机端可开始、暂停、分批报工；数量不能超过可报数量。"),
    ("质量异常", "可记录合格、不良、报废、返工及原因；数据可追溯。"),
    ("补产", "发生报废后，可形成准确的损耗记录和补产流程，不虚增库存。"),
    ("计件计算", "指定样例下，系统计件金额与双方确认的手工计算结果一致。"),
    ("权限", "工人只看到本人任务和个人计件，不能查看他人工资或管理数据。"),
    ("报表", "约定报表数据与业务单据一致，可导出Excel。"),
    ("手机访问", "客户现场Wi-Fi环境下，指定iPhone/Android浏览器可正常使用核心页面。"),
]
add_table(doc, ["验收场景", "通过标准"], acceptance_rows, [2100, 7260], font_size=9.2)
add_body(
    doc,
    "建议验收机制：客户在收到UAT版本后5个工作日内完成测试并提交问题清单。阻断核心流程的问题修复后复测；"
    "不影响核心流程的优化建议进入迭代清单，不无限延长一期验收。",
)

add_heading(doc, "11. 需求变更与商务节点建议", 1)
for text in [
    "范围冻结：第1周结束前签字/邮件确认一期清单、流程和验收案例。",
    "变更处理：新增模块、复杂规则、第三方接口、原生APP/小程序、离线能力均先评估工期和费用，再排期。",
    "付款节点建议：启动30%，核心流程演示30%，UAT版本30%，验收上线10%；最终比例由双方商务协商。",
    "维护边界建议：上线后提供约定期限的缺陷修复；新增功能、规则变化、上游大版本升级和现场服务另行报价。",
    "数据归属：客户业务数据归客户所有；实施方不得擅自使用或外泄。备份恢复责任和保留周期需在部署方案中明确。",
    "开源合规：ERPNext/Frappe基于开源许可证，核心软件和定制代码的交付、分发及品牌使用需遵守对应许可证和商标政策。",
]:
    add_bullet(doc, num_id, text)

add_heading(doc, "12. 第一次范围确认会必须回答的问题", 1)
questions = [
    "订单提交后希望系统自动到哪一步：仅提示、生成草稿，还是自动提交生产工单？",
    "企业以销定产、备货生产，还是两种模式并存？哪些产品允许直接用库存发货？",
    "缺料时是否允许部分工序先开工？“齐套”的判断按整单、工序还是批次数量？",
    "派工由计划员、班组长还是系统自动完成？是否允许多人共同完成同一工序？",
    "工人按个人计件还是班组计件？一期是否确认先做个人计件？",
    "计件数量按报工合格数还是质检最终合格数？抽检时如何折算？",
    "返工由原员工还是新员工完成？返工是否计件，责任员工是否扣减？",
    "报废扣款的合法制度依据、审批人和上限是什么？",
    "计件单价何时生效？历史工单按报工日期、工单日期还是结算日期取价？",
    "是否存在计件单价保密要求？班组长能看到哪些员工数据？",
    "基本工资、补贴、考勤扣款是否在一期一起算，还是一期只输出计件工资明细？",
    "当前考勤来自哪种设备/系统，能否提供Excel或API样例？",
    "车间网络是否稳定？是否确实存在必须离线报工的场景？",
    "一期试点选择哪1-2个产品、多少道工序、多少员工？",
    "客户指定的业务负责人、验收负责人和最终决策人分别是谁？",
]
for text in questions:
    add_bullet(doc, num_id, text)

add_heading(doc, "13. 建议的合作结论", 1)
add_callout(
    doc,
    "建议确认的合作口径",
    "以8周为一期，交付“制造主流程简化 + 手机工序报工 + 基础质量联动 + 个人计件工资”闭环。"
    "客户原始需求文档作为长期需求池，不直接作为一期全部验收依据。第1周完成范围冻结，先用1-2个真实产品跑通，"
    "再逐步扩展到更多产品、复杂计件、考勤、财务、小程序和离线能力。",
    GREEN,
)
add_body(
    doc,
    "本稿建议在下一次会议中逐条确认，并据此形成最终《项目实施范围说明书》《报价单/合同附件》和《验收用例》。",
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
